from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt

from crypto_investigator.reports.errors import DocxExportError
from crypto_investigator.reports.models import ReportDocument
from crypto_investigator.reports.typography import (
    ScriptRole,
    font_family,
    mixed_script_runs,
)


class DocxReportExporter:
    @staticmethod
    def _report_subtitle(document: ReportDocument) -> str:
        return (
            "TRX Sub-Asset Analysis and Counterparty Overview"
            if document.metadata.principal_asset_coverage == "missing"
            else "Address Profile and First-Hop Fund Flow Analysis"
        )

    @staticmethod
    def _repeat_table_header(row) -> None:
        properties = row._tr.get_or_add_trPr()
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        properties.append(header)

    @staticmethod
    def _prevent_row_split(row) -> None:
        properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        properties.append(cant_split)

    @staticmethod
    def _column_weights(columns):
        values = []
        for column in columns:
            if column in {"Address ID", "地址編號"}:
                values.append(1.05)
            elif column == "完整地址":
                values.append(3.2)
            elif column in {"排名", "信心", "方向", "資產", "完整度", "截斷"}:
                values.append(0.7)
            elif any(item in column for item in ("時間", "首次", "最後", "開始", "結束")):
                values.append(1.6)
            elif any(item in column for item in ("金額", "流入", "流出")):
                values.append(1.4)
            elif any(item in column.casefold() for item in ("地址", "來源", "去向", "sha-256")):
                values.append(1.8)
            elif any(item in column for item in ("限制", "原因", "警告", "備註", "觀察", "事實")):
                values.append(2.1)
            else:
                values.append(1.0)
        return values
    @staticmethod
    def _set_run_font(
        run, *, address: bool = False, role: ScriptRole | None = None, table=False
    ) -> None:
        family = (
            font_family(role, table=table)
            if role is not None
            else "Consolas"
            if address
            else "Times New Roman"
        )
        run.font.name = family
        run._element.get_or_add_rPr().rFonts.set(
            qn("w:eastAsia"), "標楷體"
        )
        run._element.get_or_add_rPr().rFonts.set(
            qn("w:ascii"), family
        )
        run._element.get_or_add_rPr().rFonts.set(
            qn("w:hAnsi"), family
        )
        run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), family)

    @classmethod
    def _replace_mixed_text(cls, paragraph, value, *, table=False) -> None:
        paragraph.text = ""
        for item in mixed_script_runs(value, table=table):
            if item.role is ScriptRole.NEWLINE:
                paragraph.add_run().add_break()
                continue
            run = paragraph.add_run(item.text)
            cls._set_run_font(run, role=item.role, table=table)

    @staticmethod
    def _add_page_number(paragraph) -> None:
        run = paragraph.add_run("第 ")
        DocxReportExporter._set_run_font(run, role=ScriptRole.CJK)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        field_properties = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        fonts.set(qn("w:cs"), "Times New Roman")
        field_properties.append(fonts)
        field.append(field_properties)
        run._r.addnext(field)
        text_run = paragraph.add_run(" 頁，共 ")
        DocxReportExporter._set_run_font(text_run, role=ScriptRole.CJK)
        total = OxmlElement("w:fldSimple")
        total.set(qn("w:instr"), "NUMPAGES")
        total_properties = OxmlElement("w:rPr")
        total_fonts = OxmlElement("w:rFonts")
        total_fonts.set(qn("w:ascii"), "Times New Roman")
        total_fonts.set(qn("w:hAnsi"), "Times New Roman")
        total_fonts.set(qn("w:cs"), "Times New Roman")
        total_properties.append(total_fonts)
        total.append(total_properties)
        paragraph._p.append(total)
        final_run = paragraph.add_run(" 頁")
        DocxReportExporter._set_run_font(final_run, role=ScriptRole.CJK)

    def write(self, document: ReportDocument, path: Path) -> Path:
        try:
            output = Document()
            section = output.sections[0]
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = section.bottom_margin = Mm(20)
            section.left_margin = section.right_margin = Mm(22)
            for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
                style = output.styles[style_name]
                style.font.name = "Times New Roman"
                style._element.get_or_add_rPr().rFonts.set(
                    qn("w:eastAsia"), "標楷體"
                )
                style._element.get_or_add_rPr().rFonts.set(
                    qn("w:ascii"), "Times New Roman"
                )
                style._element.get_or_add_rPr().rFonts.set(
                    qn("w:hAnsi"), "Times New Roman"
                )
            output.styles["Normal"].font.size = Pt(10.5)
            output.styles["Normal"].paragraph_format.line_spacing = 1.5
            for style_name in ("Title", "Heading 1", "Heading 2"):
                output.styles[style_name].paragraph_format.keep_with_next = True
            output.styles["Normal"].paragraph_format.widow_control = True
            self._replace_mixed_text(
                section.header.paragraphs[0],
                f"ChainSherlock | {document.metadata.report_id} | UTC+8",
            )
            footer = section.footer.paragraphs[0]
            self._add_page_number(footer)
            for report_section in document.sections:
                if report_section.section_id == "cover":
                    output.add_paragraph()
                    output.add_paragraph()
                    title = output.add_paragraph()
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._replace_mixed_text(
                        title, f"ChainSherlock {document.title}"
                    )
                    for run in title.runs:
                        run.bold = True
                        run.font.size = Pt(26)
                    subtitle = output.add_paragraph()
                    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._replace_mixed_text(
                        subtitle, self._report_subtitle(document)
                    )
                    for run in subtitle.runs:
                        run.font.size = Pt(15)
                    output.add_paragraph()
                    for block in report_section.content_blocks:
                        paragraph = output.add_paragraph()
                        self._replace_mixed_text(paragraph, block)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.line_spacing = 1.5
                        paragraph.paragraph_format.space_after = Pt(6)
                    output.add_page_break()
                    continue
                if report_section.section_id == "table_of_contents":
                    title = output.add_paragraph()
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._replace_mixed_text(title, "目錄")
                    for run in title.runs:
                        run.bold = True
                        run.font.size = Pt(24)
                    subtitle = output.add_paragraph()
                    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._replace_mixed_text(subtitle, "CONTENTS")
                    for run in subtitle.runs:
                        run.font.size = Pt(11)
                    midpoint = (len(report_section.content_blocks) + 1) // 2
                    left = report_section.content_blocks[:midpoint]
                    right = report_section.content_blocks[midpoint:]
                    toc = output.add_table(
                        rows=max(len(left), len(right)), cols=2
                    )
                    toc.autofit = False
                    for index, row in enumerate(toc.rows):
                        for column, values in enumerate((left, right)):
                            row.cells[column].width = Mm(80)
                            if index < len(values):
                                self._replace_mixed_text(
                                    row.cells[column].paragraphs[0],
                                    values[index],
                                    table=True,
                                )
                    output.add_page_break()
                    continue
                if report_section.section_id.startswith("asset_analysis_"):
                    output.add_page_break()
                if report_section.section_id == "key_addresses":
                    output.add_page_break()
                if report_section.section_id == "first_hop_candidates":
                    output.add_page_break()
                heading = output.add_heading("", level=1)
                self._replace_mixed_text(heading, report_section.title)
                heading.paragraph_format.keep_with_next = True
                for block in report_section.content_blocks:
                    paragraph = output.add_paragraph()
                    self._replace_mixed_text(paragraph, block)
                    paragraph.paragraph_format.widow_control = True
                    if report_section.section_id.startswith("ai_"):
                        paragraph.paragraph_format.keep_together = True
                for figure in report_section.figures:
                    figure_path = path.parent / figure.path
                    if figure_path.suffix.lower() not in {
                        ".png",
                        ".jpg",
                        ".jpeg",
                    }:
                        continue
                    paragraph = output.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(
                        str(figure_path),
                        width=Mm(165),
                    )
                    paragraph.paragraph_format.keep_together = True
                for table_data in report_section.tables:
                    if (
                        report_section.section_id == "key_addresses"
                        and len(table_data.columns) == 8
                    ):
                        table_heading = output.add_heading("", level=2)
                        self._replace_mixed_text(table_heading, table_data.title)
                        table_heading.paragraph_format.keep_with_next = True
                        for row in table_data.rows:
                            card = output.add_table(rows=5, cols=4)
                            card.style = "Table Grid"
                            card.autofit = False
                            values = (
                                ("調查角色", row[0], "追蹤優先級", row[6]),
                                ("完整地址（地址編號）", row[1], "", ""),
                                ("資產", row[2], "流入／流出金額", row[3]),
                                ("交易次數", row[4], "標籤狀態", row[5]),
                                ("優先理由", row[7], "", ""),
                            )
                            for row_index, values_row in enumerate(values):
                                self._prevent_row_split(card.rows[row_index])
                                for column_index, value in enumerate(values_row):
                                    self._replace_mixed_text(
                                        card.cell(
                                            row_index, column_index
                                        ).paragraphs[0],
                                        value,
                                        table=True,
                                    )
                            card.cell(1, 1).merge(card.cell(1, 3))
                            card.cell(4, 1).merge(card.cell(4, 3))
                            output.add_paragraph()
                        continue
                    if (
                        report_section.section_id == "first_hop_candidates"
                        and len(table_data.columns) == 8
                    ):
                        table_heading = output.add_heading("", level=2)
                        self._replace_mixed_text(table_heading, table_data.title)
                        table_heading.paragraph_format.keep_with_next = True
                        for row in table_data.rows:
                            card = output.add_table(rows=5, cols=4)
                            card.style = "Table Grid"
                            card.autofit = False
                            values = (
                                ("排名", row[0], "優先級", row[7]),
                                ("完整地址（地址編號）", row[1], "", ""),
                                ("收受 USDT", row[2], "占流出", row[4]),
                                ("交易次數", row[3], "標籤", row[5]),
                                ("後續資料", row[6], "", ""),
                            )
                            for row_index, values_row in enumerate(values):
                                self._prevent_row_split(card.rows[row_index])
                                for column_index, value in enumerate(values_row):
                                    cell = card.cell(row_index, column_index)
                                    self._replace_mixed_text(
                                        cell.paragraphs[0], value, table=True
                                    )
                                    for paragraph in cell.paragraphs:
                                        paragraph.paragraph_format.keep_with_next = (
                                            row_index < len(values) - 1
                                        )
                            card.cell(1, 1).merge(card.cell(1, 3))
                            card.cell(4, 1).merge(card.cell(4, 3))
                            output.add_paragraph()
                        continue
                    wide = len(table_data.columns) > 8
                    if wide:
                        current = output.add_section(WD_SECTION.NEW_PAGE)
                        current.page_width = Mm(297)
                        current.page_height = Mm(210)
                        current.top_margin = current.bottom_margin = Mm(20)
                        current.left_margin = current.right_margin = Mm(22)
                    table_heading = output.add_heading("", level=2)
                    self._replace_mixed_text(table_heading, table_data.title)
                    table_heading.paragraph_format.keep_with_next = True
                    columns = table_data.columns
                    rows = table_data.rows
                    table = output.add_table(
                        rows=1, cols=max(1, len(columns))
                    )
                    table.style = "Table Grid"
                    table.autofit = False
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    weights = (
                        (0.6, 4.0, 1.5, 1.0, 1.2, 0.8)
                        if table_data.table_id
                        == "first_hop_candidates_flow"
                        else self._column_weights(columns)
                    )
                    usable_mm = 253 if wide else 166
                    total_weight = sum(weights) or 1
                    for index, column in enumerate(columns):
                        width = Mm(usable_mm * weights[index] / total_weight)
                        table.columns[index].width = width
                        self._replace_mixed_text(
                            table.rows[0].cells[index].paragraphs[0],
                            column,
                            table=True,
                        )
                        table.rows[0].cells[index].width = width
                        table.rows[0].cells[index].vertical_alignment = (
                            WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        )
                    self._repeat_table_header(table.rows[0])
                    for row in rows:
                        cells = table.add_row().cells
                        self._prevent_row_split(table.rows[-1])
                        if table_data.table_id == "key_address_summary":
                            table.rows[-1].height = Mm(18)
                            table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                        for index, value in enumerate(row):
                            self._replace_mixed_text(
                                cells[index].paragraphs[0], value, table=True
                            )
                            cells[index].width = Mm(
                                usable_mm * weights[index] / total_weight
                            )
                            cells[index].vertical_alignment = (
                                WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            )
                            if any(
                                marker in columns[index]
                                for marker in (
                                    "金額", "筆數", "交易數", "比例",
                                    "占比", "事件數", "取得筆數",
                                )
                            ):
                                cells[index].paragraphs[0].alignment = (
                                    WD_ALIGN_PARAGRAPH.RIGHT
                                )
                    if wide:
                        current = output.add_section(WD_SECTION.NEW_PAGE)
                        current.page_width = Mm(210)
                        current.page_height = Mm(297)
                        current.top_margin = current.bottom_margin = Mm(20)
                        current.left_margin = current.right_margin = Mm(22)
                if report_section.section_id == "key_addresses":
                    output.add_page_break()
                if report_section.section_id == "first_hop_candidates":
                    output.add_page_break()
                if report_section.section_id == "table_of_contents":
                    output.add_page_break()
            path.parent.mkdir(parents=True, exist_ok=True)
            output.save(path)
            return path
        except Exception as error:
            raise DocxExportError("Unable to export DOCX report") from error
