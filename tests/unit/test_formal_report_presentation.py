from dataclasses import replace
from copy import deepcopy
from datetime import UTC, datetime
from decimal import Decimal
from html import unescape
import re
from types import SimpleNamespace
import zipfile

import pytest
from docx import Document

from crypto_investigator.cli import _provider_scope
from crypto_investigator.domain.scope import PaginationPolicy, ScopeType
from crypto_investigator.reports.docx_exporter import DocxReportExporter
from crypto_investigator.reports.export import ReportExportCoordinator
from crypto_investigator.reports.html_exporter import HtmlReportExporter
from crypto_investigator.reports.composer import ReportComposer
from crypto_investigator.reports.formatting import (
    abbreviate_identifier,
    format_amount,
    format_duration,
    format_percent,
)
from crypto_investigator.reports.models import (
    ReportEvidence,
    ReportLimitation,
    ReportSection,
    ReportTable,
)
from crypto_investigator.reports.forensic_artifacts import (
    suspicious_trx_candidates,
    trx_reconciliation,
)
from crypto_investigator.reports.pdf_exporter import PdfReportExporter
from crypto_investigator.reports.presentation import (
    format_display_text,
    prepare_report_for_display,
)


ADDRESS = "TR5WMAhpM9JkpouAT49X9pNHP8NPQkcGAE"
OTHER = "TGxVDjGujrUXQGZtWgNcdeawkPYeoE4Frv"
THIRD = "TJMUadmxstaJsnsK6my4vEDGCDiXMh3eWd"


def _analysis():
    return {
        "summary": {
            "first_seen": "2026-01-05T03:31:27+00:00",
            "last_seen": "2026-01-06T03:31:27+00:00",
            "transaction_count": 1000,
            "incoming_count": 415,
            "outgoing_count": 477,
            "unique_counterparties": 12,
        },
        "statistics": {
            "incoming_amount": {"USDT": Decimal("1234567.89"), "0597 COM": Decimal("0.01")},
            "outgoing_amount": {"USDT": Decimal("1200000"), "0597 COM": Decimal("0")},
            "asset_breakdown": {
                "USDT": {"transaction_count": 999},
                "0597 COM": {"transaction_count": 1},
            },
        },
        "counterparties": [{
            "address": OTHER,
            "incoming_count": 2,
            "outgoing_count": 8,
            "interaction_count": 10,
            "incoming_amount_by_asset": {"USDT": Decimal("100")},
            "outgoing_amount_by_asset": {"USDT": Decimal("900")},
            "first_seen": "2026-01-05T03:31:27+00:00",
            "last_seen": "2026-01-06T03:31:27+00:00",
        }],
        "timeline": {},
        "metadata": {
            "completeness": "partial",
            "analysis_scope": {
                "scope_type": "full_history",
                "timezone": "Asia/Taipei",
            },
            "time_scope": {"full_history_complete": False},
        },
        "warnings": [],
    }


def _investigation():
    return {
        "structured_metadata": {
            "source_transaction_count": 1000,
            "assets": ["USDT"],
        },
        "direction_reconciliation": {
            "failed_transaction_count": 0,
            "unclassified_direction_count": 108,
        },
        "funding": {
            "sources": [{
                "address": ADDRESS,
                "amounts_by_asset": {"USDT": "291166.569"},
                "share_by_asset": {"USDT": "0.291166569"},
                "first_funding": "2026-01-05T03:31:27+00:00",
                "last_funding": "2026-01-06T03:31:27+00:00",
            }],
            "transitions": [],
            "top_sources_by_asset": {"USDT": [ADDRESS]},
            "concentration_by_asset": {},
        },
        "distribution_analysis": {"statistics_by_asset": {
            "USDT": {
                "matched_incoming_amount": "100",
                "matched_outgoing_amount": "90",
                "unmatched_incoming_amount": "10",
                "unmatched_outgoing_amount": "5",
                "average_holding_seconds": "86254.5",
                "median_holding_seconds": "248397",
                "within_1_hour_ratio": "0.1",
                "within_24_hours_ratio": "0.75",
                "pass_through_event_count": 8,
            }
        }},
        "stages": [{
            "stage": "dominant",
            "started_at": "2026-01-05T03:31:27+00:00",
            "ended_at": "2026-01-06T03:31:27+00:00",
            "transaction_count": 100,
            "assets": ["USDT"],
            "dominant_funding_sources": [ADDRESS],
            "dominant_outgoing_counterparties": [OTHER],
            "reason_codes": ["frequency_increased"],
        }],
        "dormant_periods": [],
        "transfer_patterns": {
            "integer_amount_ratio": "0.75",
            "batch_incoming_count": 4,
            "batch_outgoing_count": 49,
            "fixed_amounts": {"USDT": [500, 1000, 2000]},
            "amount_suffix_counts": {"00": 400},
        },
        "services": [],
        "observations": [{
            "code": "funding_source_changed",
            "factual_statement": "主要供款方於 2026-01-05T03:31:27+00:00 改變。",
            "evidence_refs": ["OBS-001"],
        }],
        "conclusion_fact_items": [{
            "fact_code": "provider_truncated",
            "value": True,
            "evidence_refs": ["FACT-003"],
        }],
        "evidence_refs": [{"evidence_id": "IF0"}],
    }


def _document():
    return ReportComposer().compose(
        _analysis(),
        investigation=_investigation(),
        target_address=ADDRESS,
        chain="tron",
        provider_status=({
            "chain": "tron",
            "capability": "token_transfers",
            "provider": "trongrid",
            "fetched_records": 1000,
            "completeness": "partial",
            "truncated": True,
            "truncation_reason": "safety_limit",
            "warnings": ["仍有更多資料"],
            "analysis_completeness": "internal",
            "fallback_attempted": True,
        },),
        materiality_thresholds={"0597 COM": Decimal("1")},
    )


def _section(document, section_id):
    return next(item for item in document.sections if item.section_id == section_id)


def test_01_provider_table_has_eight_columns():
    assert len(_section(_document(), "provider_status").tables[0].columns) == 8


def test_02_provider_engineering_columns_are_absent():
    columns = _section(_document(), "provider_status").tables[0].columns
    assert "analysis_completeness" not in columns
    assert "fallback_attempted" not in columns


def test_03_counterparty_table_has_no_duplicate_count_columns():
    columns = _section(_document(), "outgoing_distribution").tables[0].columns
    assert columns.count("交易次數") == 1
    assert "incoming_count" not in columns


def test_04_main_ranking_uses_registry_id_without_duplicate_address():
    display = prepare_report_for_display(_document())
    table = next(
        table
        for table in _section(display, "asset_analysis_usdt").tables
        if table.table_id == "outgoing_rank_usdt"
    )
    row = table.rows[0]
    assert row[1].startswith("地址-")
    assert "地址" not in table.columns
    assert OTHER not in row


def test_05_full_address_is_preserved_in_front_address_index():
    display = prepare_report_for_display(_document())
    table = next(
        table
        for table in _section(display, "address_registry").tables
        if table.table_id == "address_registry_identity"
    )


def _two_asset_document():
    analysis = deepcopy(_analysis())
    analysis["statistics"]["incoming_amount"]["TRX"] = Decimal("5000")
    analysis["statistics"]["outgoing_amount"]["TRX"] = Decimal("4112")
    analysis["statistics"]["asset_breakdown"]["TRX"] = {"transaction_count": 684}
    analysis["counterparties"].append({
        "address": THIRD,
        "incoming_count": 0,
        "outgoing_count": 684,
        "interaction_count": 684,
        "incoming_amount_by_asset": {"TRX": Decimal("0")},
        "outgoing_amount_by_asset": {"TRX": Decimal("4112")},
        "first_seen": "2025-06-02T09:21:48+00:00",
        "last_seen": "2026-08-04T10:24:57+00:00",
    })
    investigation = deepcopy(_investigation())
    investigation["structured_metadata"]["assets"].append("TRX")
    investigation["funding"]["sources"].append({
        "address": THIRD,
        "amounts_by_asset": {"TRX": "5000"},
        "share_by_asset": {"TRX": "0.75"},
        "first_funding": "2025-06-02T09:21:48+00:00",
        "last_funding": "2026-08-04T10:24:57+00:00",
    })
    investigation["funding"]["top_sources_by_asset"]["TRX"] = [THIRD]
    return ReportComposer().compose(
        analysis,
        investigation=investigation,
        target_address=ADDRESS,
        chain="tron",
    )
    assert any(OTHER in row for row in table.rows)


def _trx_only_document():
    analysis = deepcopy(_analysis())
    analysis["statistics"] = {
        "incoming_amount": {"TRX": Decimal("5243.21")},
        "outgoing_amount": {"TRX": Decimal("4621.52")},
        "asset_breakdown": {"TRX": {"transaction_count": 3034}},
    }
    analysis["counterparties"] = [{
        "address": THIRD,
        "incoming_count": 0,
        "outgoing_count": 684,
        "interaction_count": 684,
        "incoming_amount_by_asset": {"TRX": Decimal("0")},
        "outgoing_amount_by_asset": {"TRX": Decimal("4112")},
        "first_seen": "2025-06-02T09:21:48+00:00",
        "last_seen": "2026-08-04T10:24:57+00:00",
    }]
    investigation = deepcopy(_investigation())
    investigation["structured_metadata"]["assets"] = ["TRX"]
    investigation["funding"]["sources"] = [{
        "address": OTHER,
        "amounts_by_asset": {"TRX": "4860.72"},
        "share_by_asset": {"TRX": "0.927"},
        "first_funding": "2025-06-02T09:21:48+00:00",
        "last_funding": "2026-08-04T10:24:57+00:00",
    }]
    investigation["funding"]["top_sources_by_asset"] = {"TRX": [OTHER]}
    return ReportComposer().compose(
        analysis,
        investigation=investigation,
        target_address=ADDRESS,
        chain="tron",
    )


def test_06_python_dict_repr_is_absent_from_main_tables():
    rendered = str([table.rows for section in _document().sections for table in section.tables])
    assert "{'" not in rendered


def test_07_python_list_repr_is_absent_from_pattern_table():
    table = _section(_document(), "transfer_patterns").tables[0]
    assert "[" not in str(table.rows)


def test_08_ratio_is_rendered_as_percentage():
    table = _section(_document(), "funding_analysis").tables[0]
    assert table.rows[0][4] == "29.12%"


def test_09_amount_uses_grouping_without_losing_precision():
    assert format_amount("1234567.89") == "1,234,567.89"


def test_10_duration_under_one_day_is_human_readable():
    assert format_duration("86254.5") == "23 小時 57 分"


def test_11_duration_over_one_day_is_human_readable():
    assert format_duration("248397") == "2 天 20 小時 59 分"


def test_12_utc_is_converted_to_taipei():
    assert format_display_text(
        "2026-01-05T03:31:27+00:00", "Asia/Taipei"
    ) == "2026-01-05 11:31:27"


def test_13_naive_datetime_is_marked_unknown():
    assert "timezone unknown" in format_display_text(
        "2026-01-05T03:31:27", "Asia/Taipei"
    )


def test_14_metadata_uses_case_timezone():
    assert _document().metadata.timezone == "Asia/Taipei"


def test_15_cover_declares_taipei_timezone():
    assert "UTC+8（Asia/Taipei）" in " ".join(
        _section(_document(), "cover").content_blocks
    )


def test_16_display_generated_at_uses_taipei_offset():
    document = replace(
        _document(),
        metadata=replace(
            _document().metadata,
            generated_at=datetime(2026, 1, 5, 3, 31, 27, tzinfo=UTC),
        ),
    )
    assert prepare_report_for_display(document).metadata.generated_at.hour == 11


def test_17_ai_enrichment_datetime_text_uses_taipei():
    document = _document()
    extra = ReportSection(
        "ai_test", "AI 專業綜合", 30,
        ("發生時間 2026-01-05T03:31:27+00:00。",),
    )
    display = prepare_report_for_display(
        replace(document, sections=(*document.sections, extra))
    )
    assert "2026-01-05 11:31:27" in _section(display, "ai_test").content_blocks[0]


def test_18_evidence_index_is_artifact_level():
    table = _section(_document(), "evidence_index").tables[0]
    assert table.columns == (
        "Evidence ID", "檔名", "類型", "SHA-256", "完整性", "來源", "備註"
    )


def test_19_record_level_if_id_is_not_in_evidence_index():
    table = _section(_document(), "evidence_index").tables[0]
    assert not any(cell.startswith("IF") for row in table.rows for cell in row)


def test_20_legacy_evidence_is_not_marked_verified():
    row = _section(_document(), "evidence_index").tables[0].rows[0]
    assert row[3] == "雜湊不可用"
    assert row[4] == "無法驗證"


def test_21_ai_section_order_precedes_evidence_index():
    document = _document()
    ai = ReportSection("ai_conclusion", "AI 專業綜合", 30, ("內容",))
    ordered = sorted((*document.sections, ai), key=lambda item: (item.order, item.section_id))
    assert [item.section_id for item in ordered].index("ai_conclusion") < [
        item.section_id for item in ordered
    ].index("evidence_index")


def test_22_display_tables_are_limited_to_ten_columns():
    document = _document()
    wide = ReportTable(
        "wide", "wide", tuple(f"c{i}" for i in range(12)),
        (tuple(str(i) for i in range(12)),),
    )
    extra = ReportSection("wide", "wide", 40, tables=(wide,))
    display = prepare_report_for_display(
        replace(document, sections=(*document.sections, extra))
    )
    assert all(
        len(table.columns) <= 6 for table in _section(display, "wide").tables
    )
    assert len(_section(display, "wide").tables) == 3


def test_23_all_composed_main_tables_have_at_most_ten_columns():
    assert max(
        len(table.columns)
        for section in _document().sections
        for table in section.tables
    ) <= 10


def test_24_non_material_asset_is_moved_to_appendix():
    assert "0597 COM" not in {
        row[0] for row in _section(_document(), "asset_flows").tables[0].rows
    }
    assert "0597 COM" in {
        row[0] for row in _section(_document(), "non_material_assets").tables[0].rows
    }


def test_25_percentage_formatter_is_stable():
    assert format_percent("0.291166569") == "29.12%"


def test_26_formal_full_history_is_unbounded():
    settings = SimpleNamespace(
        pagination=SimpleNamespace(max_pages=1, max_records=50)
    )
    scope = _provider_scope(
        ScopeType.FULL_HISTORY, None, None,
        timezone="Asia/Taipei", settings=settings,
    )
    assert scope.pagination_policy is PaginationPolicy.TO_PROVIDER_END
    assert scope.max_pages is None
    assert scope.max_records is None


def test_27_quick_preview_is_explicitly_bounded():
    settings = SimpleNamespace(
        pagination=SimpleNamespace(max_pages=1, max_records=500)
    )
    scope = _provider_scope(
        ScopeType.QUICK_PREVIEW, None, None,
        timezone="Asia/Taipei", settings=settings,
    )
    assert scope.pagination_policy is PaginationPolicy.BOUNDED
    assert scope.max_pages == 1
    assert scope.max_records == 500


def test_28_full_history_rejects_preview_limits():
    settings = SimpleNamespace(
        pagination=SimpleNamespace(max_pages=1, max_records=500)
    )
    with pytest.raises(Exception):
        _provider_scope(
            ScopeType.FULL_HISTORY, 1, 1000,
            timezone="Asia/Taipei", settings=settings,
        )


def test_29_complete_full_history_uses_address_first_seen_label():
    value = _analysis()
    value["metadata"]["completeness"] = "complete"
    value["metadata"]["time_scope"]["full_history_complete"] = True
    document = ReportComposer().compose(value, investigation=_investigation())
    rows = _section(document, "analysis_summary").tables[0].rows
    assert any(row[0] == "地址首次交易時間" for row in rows)


def test_30_asset_time_scope_is_preserved_by_asset():
    value = _analysis()
    value["statistics"]["incoming_amount"]["TRX"] = Decimal("10")
    value["statistics"]["outgoing_amount"]["TRX"] = Decimal("5")
    value["statistics"]["asset_breakdown"]["TRX"] = {"transaction_count": 2}
    value["metadata"]["time_scope"].update({
        "first_seen_by_asset": {
            "TRX": "2025-05-30T16:47:06Z",
            "USDT": "2025-05-30T16:55:36Z",
        },
        "last_seen_by_asset": {
            "TRX": "2026-08-04T15:01:21Z",
            "USDT": "2026-08-04T14:59:15Z",
        },
    })
    table = _section(
        ReportComposer().compose(value, investigation=_investigation()),
        "asset_flows",
    ).tables[1]
    assert ("TRX", "2025-05-31 00:47:06", "2026-08-04 23:01:21") in table.rows
    assert ("USDT", "2025-05-31 00:55:36", "2026-08-04 22:59:15") in table.rows


def test_31_docx_has_repeating_headers_and_non_splitting_rows(tmp_path):
    path = DocxReportExporter().write(
        prepare_report_for_display(_document()), tmp_path / "report.docx"
    )
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:tblHeader" in xml
    assert "<w:cantSplit" in xml
    assert "<w:keepNext" in xml


def test_32_html_print_layout_repeats_headers_and_keeps_rows(tmp_path):
    path = HtmlReportExporter().write(
        prepare_report_for_display(_document()), tmp_path / "report.html"
    )
    content = path.read_text(encoding="utf-8")
    assert "<thead>" in content
    assert "display:table-header-group" in content
    assert "page-break-inside:avoid" in content


def test_33_evidence_index_remains_artifact_level_and_compact():
    table = _section(_document(), "evidence_index").tables[0]
    assert len(table.rows) < 10
    assert all(not row[0].startswith("IF") for row in table.rows)


def test_34_booklet_starts_with_cover_then_contents():
    ids = [item.section_id for item in prepare_report_for_display(_document()).sections]
    assert ids[:2] == ["cover", "table_of_contents"]


def test_35_asset_analysis_precedes_address_rankings():
    ids = [item.section_id for item in prepare_report_for_display(_two_asset_document()).sections]
    assert ids.index("asset_analysis_usdt") < ids.index("address_rankings")
    assert ids.index("asset_analysis_trx") < ids.index("address_rankings")


def test_36_usdt_and_trx_are_separate_chapters():
    display = prepare_report_for_display(_two_asset_document())
    assert (
        _section(display, "asset_analysis_usdt").title
        == "USDT 主要價值資產分析"
    )
    assert (
        _section(display, "asset_analysis_trx").title
        == "TRX 營運資產與費用型對手方分析"
    )


@pytest.mark.parametrize("asset", ("usdt", "trx"))
def test_37_each_asset_has_three_distinct_rankings(asset):
    table_ids = {
        table.table_id
        for table in _section(
            prepare_report_for_display(_two_asset_document()),
            f"asset_analysis_{asset}",
        ).tables
    }
    assert {
        f"funding_rank_{asset}",
        f"outgoing_rank_{asset}",
        f"frequency_rank_{asset}",
    } <= table_ids


def test_38_ranking_titles_state_the_basis():
    display = prepare_report_for_display(_two_asset_document())
    titles = [
        table.title
        for asset in ("usdt", "trx")
        for table in _section(display, f"asset_analysis_{asset}").tables
    ]
    assert any("依流入金額" in title for title in titles)
    assert any("依流出金額" in title for title in titles)
    assert any("依交易次數" in title for title in titles)


def test_39_ranking_tables_never_mix_assets():
    display = prepare_report_for_display(_two_asset_document())
    for asset in ("USDT", "TRX"):
        section = _section(display, f"asset_analysis_{asset.casefold()}")
        assert all(
            other not in table.title
            for table in section.tables
            for other in ({"USDT", "TRX"} - {asset})
        )


def test_40_main_ranking_uses_address_id_only():
    table = next(
        table
        for table in _section(
            prepare_report_for_display(_two_asset_document()),
            "asset_analysis_trx",
        ).tables
        if table.table_id == "outgoing_rank_trx"
    )
    assert table.rows[0][1].startswith("地址-")
    assert THIRD not in table.rows[0]


def test_41_address_ids_are_stable_across_sections():
    display = prepare_report_for_display(_two_asset_document())
    registry = next(
        table for table in _section(display, "address_registry").tables
        if table.table_id == "address_registry_identity"
    )
    known = {row[2]: row[0] for row in registry.rows}
    key = next(
        table for table in _section(display, "key_addresses").tables
        if table.table_id == "key_address_summary"
    )


def _repeated_promotional_document():
    document = _two_asset_document()
    sections = []
    senders = (
        "TQPtmCQeYzn1iUWv6sun2aoKRYiBrB4Aq4",
        "TL492pHAGYppvE8QKNwai8GakhuWjB8uE7",
    )
    for section in document.sections:
        if section.section_id == "funding_analysis":
            tables = []
            for table in section.tables:
                if table.table_id == "funding_sources":
                    rows = tuple(table.rows) + tuple(
                        (
                            str(index),
                            "TRX",
                            sender,
                            "8888.88",
                            "10%",
                            "2025-06-01T00:00:00+00:00",
                            "2025-06-01T00:00:00+00:00",
                        )
                        for index, sender in enumerate(senders, 20)
                    )
                    table = replace(table, rows=rows)
                tables.append(table)
            section = replace(section, tables=tuple(tables))
        sections.append(section)
    return replace(document, sections=tuple(sections))
    for row in key.rows:
        full_address, display_id = row[1].rsplit("（", 1)
        assert display_id.rstrip("）") == known[full_address]


def test_41a_key_addresses_start_with_contiguous_priority_ids():
    display = prepare_report_for_display(_two_asset_document())
    reference = next(
        table for table in _section(display, "address_registry").tables
        if table.table_id == "address_registry_identity"
    )
    unique_ids = list(dict.fromkeys(row[0] for row in reference.rows))
    assert unique_ids[0] == "地址-001"
    assert unique_ids == sorted(unique_ids)


def test_41b_booklet_amounts_round_half_up_to_two_decimals():
    table = next(
        table
        for table in _section(
            prepare_report_for_display(_document()),
            "asset_analysis_usdt",
        ).tables
        if table.table_id == "funding_rank_usdt"
    )
    assert table.rows[0][2] == "291,166.57"


def test_41c_every_main_ranking_id_maps_to_a_full_registry_address():
    display = prepare_report_for_display(_two_asset_document())
    reference = next(
        table
        for table in _section(display, "address_registry").tables
        if table.table_id == "address_registry_identity"
    )
    mapped = {row[0] for row in reference.rows}
    used = {
        row[1].splitlines()[0]
        for section in display.sections
        if section.section_id.startswith("asset_analysis_")
        for table in section.tables
        if table.table_id.startswith(
            ("funding_rank_", "outgoing_rank_", "frequency_rank_")
        )
        for row in table.rows
    }
    assert used <= mapped
    assert all(len(row[2]) > 20 for row in reference.rows)
    numbers = [int(row[0].split("-")[1]) for row in reference.rows]
    assert numbers == list(range(1, len(numbers) + 1))


def test_41d_key_address_summary_contains_complete_copy_safe_addresses():
    key = next(
        table
        for table in _section(
            prepare_report_for_display(_two_asset_document()),
            "key_addresses",
        ).tables
        if table.table_id == "key_address_summary"
    )
    values = {row[1] for row in key.rows}
    assert f"{ADDRESS}（地址-001）" in values
    assert all("…" not in value for value in values)
    assert all(value.split("（", 1)[0].strip() == value.split("（", 1)[0] for value in values)


def test_41e_narrative_uses_one_consistent_abbreviation_with_display_id():
    display = prepare_report_for_display(_two_asset_document())
    text = "\n".join(
        block
        for section in display.sections
        for block in section.content_blocks
    )
    expected = f"{abbreviate_identifier(ADDRESS)}（地址-001）"
    assert expected in text
    assert ADDRESS not in "\n".join(
        block
        for section in display.sections
        if section.section_id != "cover"
        for block in section.content_blocks
    )


def test_41f_ranking_table_uses_id_and_fixed_abbreviation_not_full_address():
    display = prepare_report_for_display(_two_asset_document())
    tables = [
        table
        for section in display.sections
        if section.section_id.startswith("asset_analysis_")
        for table in section.tables
        if table.table_id.startswith(
            ("funding_rank_", "outgoing_rank_", "frequency_rank_")
        )
    ]
    references = [row[1] for table in tables for row in table.rows]
    assert all(value.startswith("地址-") and "\n" in value for value in references)
    assert all(ADDRESS not in value for value in references)
    assert all(
        value.splitlines()[1] == abbreviate_identifier(next(
            row[2] for row in next(
                table for table in _section(display, "address_registry").tables
                if table.table_id == "address_registry_identity"
            ).rows if row[0] == value.splitlines()[0]
        ))
        for value in references
    )


def test_42_key_address_summary_is_early():
    ids = [item.section_id for item in prepare_report_for_display(_two_asset_document()).sections]
    assert ids.index("key_addresses") < ids.index("asset_analysis_usdt")


def test_43_fund_flow_path_section_exists():
    section = _section(prepare_report_for_display(_two_asset_document()), "fund_flow_paths")
    assert section.tables[0].rows
    assert section.title == "主要來源與去向關聯摘要"
    assert all(row[2] == "候選摘要" for row in section.tables[1].rows)
    assert all("不代表同一筆資金" in row[4] for row in section.tables[1].rows)


def test_44_ai_sections_follow_deterministic_facts():
    document = prepare_report_for_display(_document())
    ids = [item.section_id for item in document.sections]
    ai = [index for index, value in enumerate(ids) if value.startswith("ai_")]
    assert not ai or ids.index("investigation_facts") < min(ai)


def test_45_address_registry_is_before_asset_analysis():
    ids = [item.section_id for item in prepare_report_for_display(_document()).sections]
    assert ids.index("key_addresses") < ids.index("address_registry")
    assert ids.index("address_registry") < ids.index("asset_flows")


def test_46_address_registry_csv_contains_complete_values(tmp_path):
    result = ReportExportCoordinator().export(
        _two_asset_document(), tmp_path, requested_format="markdown"
    )
    content = (tmp_path / result.files["address_registry"]).read_text(
        encoding="utf-8-sig"
    )
    assert ADDRESS in content
    assert THIRD in content


def test_47_cover_and_contents_are_forced_to_separate_docx_pages(tmp_path):
    path = DocxReportExporter().write(
        prepare_report_for_display(_document()), tmp_path / "report.docx"
    )
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert xml.count('<w:br w:type="page"') >= 2


def test_48_html_has_booklet_page_breaks(tmp_path):
    path = HtmlReportExporter().write(
        prepare_report_for_display(_document()), tmp_path / "report.html"
    )
    content = path.read_text(encoding="utf-8")
    assert "#cover" in content and "#table_of_contents" in content
    assert "page-break-after:always" in content


def test_49_four_format_address_display_mapping_is_consistent(tmp_path):
    display = prepare_report_for_display(_two_asset_document())
    full = f"{ADDRESS}（地址-001）"
    compact = f"地址-001\n{abbreviate_identifier(ADDRESS)}"

    html_path = HtmlReportExporter().write(display, tmp_path / "report.html")
    html = html_path.read_text(encoding="utf-8")
    html_text = unescape(
        re.sub(r"<[^>]+>", "", re.sub(r"<br\s*/?>", "\n", html))
    )
    assert full in html_text
    assert compact in html_text

    markdown_path = ReportExportCoordinator().export(
        _two_asset_document(), tmp_path / "markdown", requested_format="markdown"
    )
    markdown = (
        tmp_path / "markdown" / markdown_path.files["markdown"]
    ).read_text(encoding="utf-8")
    assert full in markdown
    assert compact.replace("\n", "<br>") in markdown

    docx_path = DocxReportExporter().write(display, tmp_path / "report.docx")
    docx = Document(docx_path)
    docx_text = "\n".join(
        [
            *(paragraph.text for paragraph in docx.paragraphs),
            *(
                cell.text
                for table in docx.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )
    assert ADDRESS in docx_text and "地址-001" in docx_text
    assert abbreviate_identifier(ADDRESS) in docx_text

    assert PdfReportExporter._pdf_cell(compact, "地址參照") == compact


def test_50_primary_tables_are_bounded():
    display = prepare_report_for_display(_two_asset_document())
    assert all(
        len(table.rows) <= 10
        for section in display.sections
        if section.section_id.startswith("asset_analysis_")
        for table in section.tables
    )


def test_50_engineering_tables_are_not_in_the_booklet_body():
    ids = {item.section_id for item in prepare_report_for_display(_document()).sections}
    assert "data_pipeline" not in ids
    assert "provider_status" not in ids
    assert "rejected_records" not in ids


def test_general_report_omits_address_pollution_section():
    display = prepare_report_for_display(_two_asset_document())
    assert "address_pollution_safety" not in {
        section.section_id for section in display.sections
    }


def test_empty_address_pollution_result_does_not_emit_empty_section():
    document = _two_asset_document()
    document = replace(
        document,
        metadata=replace(
            document.metadata,
            first_hop_product={"address_pollution": None},
        ),
    )
    display = prepare_report_for_display(document)
    assert "address_pollution_safety" not in {
        section.section_id for section in display.sections
    }


def test_51_asset_chapter_begins_with_narrative():
    section = _section(
        prepare_report_for_display(_two_asset_document()),
        "asset_analysis_usdt",
    )
    assert "共納入" in section.content_blocks[0]
    assert "最大資金來源" in " ".join(section.content_blocks)


def test_52_confirmed_facts_are_direct_data_not_rule_results():
    display = prepare_report_for_display(_document())
    section = _section(display, "confirmed_facts")
    rendered = str(section.tables[0].rows)
    assert "FACT-COUNT-001" in rendered
    assert "有辨識到" not in rendered
    assert "未辨識到" not in rendered
    assert "batch" not in rendered.casefold()
    assert "investigation_facts" not in {
        item.section_id for item in display.sections
    }


def test_53_candidate_flow_is_not_presented_as_traced_path():
    section = _section(
        prepare_report_for_display(_two_asset_document()),
        "fund_flow_paths",
    )
    rendered = f"{section.title} {section.tables}"
    assert "主要資金流路徑" not in rendered
    assert "不代表同一筆資金" in rendered
    assert "來源與去向排名組合" in rendered


def test_54_aggregate_amounts_are_not_used_to_infer_trx_candidates():
    document = _repeated_promotional_document()
    candidates = suspicious_trx_candidates(document)
    assert candidates == ()
    reconciliation = trx_reconciliation(document)
    assert Decimal(reconciliation["gross_on_chain_inflow"]) == Decimal("5000")
    assert Decimal(reconciliation["promotional_candidate"]) == Decimal("0")
    assert Decimal(reconciliation["final_material_inflow"]) == Decimal("5000")


def test_55_report_does_not_create_suspicious_trx_table_from_amounts():
    display = prepare_report_for_display(_repeated_promotional_document())
    trx = _section(display, "asset_analysis_trx")
    assert all(
        not table.table_id.startswith("suspicious_trx")
        for table in trx.tables
    )


def test_56_operation_stage_uses_chinese_confidence_and_bounded_assets():
    display = prepare_report_for_display(_document())
    stages = _section(display, "operation_stages")
    assert stages.tables
    rendered = str(stages.tables)
    assert "medium" not in rendered
    assert any(value in rendered for value in ("低", "中", "高"))
    assert all(len(table.columns) == 2 for table in stages.tables)


def test_57_graph_truncation_contradiction_is_overridden_without_ai_call():
    document = _document()
    contradiction = ReportSection(
        "ai_conclusion",
        "AI 專業綜合",
        500,
        ("圖譜與供應端均未標示截斷。",),
    )
    document = replace(
        document,
        sections=(*document.sections, contradiction),
        limitations=(
            *document.limitations,
            ReportLimitation("graph_truncated", "安全上限"),
        ),
    )
    display = prepare_report_for_display(document)
    rendered = " ".join(_section(display, "ai_conclusion").content_blocks)
    assert "圖譜與供應端均未標示截斷" not in rendered
    assert "Graph 因安全上限截斷" in rendered


def test_58_display_does_not_use_if0_as_primary_claim_reference():
    display = prepare_report_for_display(_document())
    rendered = str(display.sections)
    assert "IF0" not in rendered


def test_59_graph_fact_uses_graph_section_as_single_source_of_truth():
    document = _document()
    sections = tuple(
        replace(section, content_blocks=("節點：17；邊：17；截斷：否。",))
        if section.section_id == "graph"
        else section
        for section in document.sections
    )
    display = prepare_report_for_display(replace(document, sections=sections))
    facts = _section(display, "confirmed_facts").tables[0]
    graph = next(row for row in facts.rows if row[0] == "FACT-GRAPH-001")
    assert graph[2] == "未標示截斷"
    assert facts.columns == ("事實編號", "事實內容", "數值")
    assert len(graph) == 3


def test_60_direction_fact_and_layered_counts_use_report_metadata():
    document = _document()
    metadata = replace(
        document.metadata,
        transaction_count=3099,
        incoming_count=2410,
        outgoing_count=689,
        unclassified_count=0,
        native_trx_transaction_count=3034,
        other_asset_transaction_count=65,
        micro_excluded_count=2316,
        analysis_record_count=718,
        retrieval_completeness="complete",
        asset_classification_completeness="complete",
        material_analysis_scope="718 records",
        graph_completeness="complete",
        graph_node_count=17,
        graph_edge_count=17,
    )
    display = prepare_report_for_display(replace(document, metadata=metadata))
    facts = _section(display, "confirmed_facts").tables[0]
    directions = {
        row[0]: row[2]
        for row in facts.rows
        if row[0].startswith("FACT-DIRECTION-")
    }
    assert directions == {
        "FACT-DIRECTION-IN-001": "2,410",
        "FACT-DIRECTION-OUT-001": "689",
        "FACT-DIRECTION-UNKNOWN-001": "0",
    }
    layers = _section(display, "completeness_layers").tables[0]
    values = {row[0]: row[1] for row in layers.rows}
    assert [
        values["完整取得交易"],
        values["原生 TRX"],
        values["原生 TRX 流入／流出"],
        values["TRC10／其他資產"],
        values["TRC10／其他資產流入"],
        values["微額 TRX 技術性排除"],
        values["主要資金流與行為分析"],
    ] == [
        "3,099", "3,034", "2,345／689", "65", "65", "2,316", "718"
    ]


def test_61_absent_usdt_does_not_create_usdt_fact():
    document = _document()
    sections = tuple(
        replace(
            section,
            tables=tuple(
                replace(
                    table,
                    rows=tuple(
                        ("TRX", *row[1:]) if table.table_id == "asset_flows" else row
                        for row in table.rows
                    ),
                )
                for table in section.tables
            ),
        )
        for section in document.sections
    )
    display = prepare_report_for_display(replace(document, sections=sections))
    facts = _section(display, "confirmed_facts").tables[0]
    assert all(row[0] != "FACT-ASSET-USDT-001" for row in facts.rows)


def test_62_funding_ranking_stops_at_material_sources():
    document = _two_asset_document()
    addresses = tuple(f"T{'A' * 32}{suffix}" for suffix in "23456")
    funding_rows = tuple(
        (
            str(index),
            "TRX",
            address,
            amount,
            share,
            "2025-01-01T00:00:00+00:00",
            "2025-02-01T00:00:00+00:00",
        )
        for index, (address, amount, share) in enumerate(
            zip(
                addresses,
                ("4668.72", "202.99", "91.66", "7", "0.64"),
                ("89.04%", "3.87%", "1.75%", "0.13%", "0.01%"),
            ),
            1,
        )
    )
    sections = tuple(
        replace(
            section,
            tables=tuple(
                replace(table, rows=funding_rows)
                if table.table_id == "funding_sources"
                else table
                for table in section.tables
            ),
        )
        for section in document.sections
    )
    display = prepare_report_for_display(replace(document, sections=sections))
    table = next(
        table
        for table in _section(display, "asset_analysis_trx").tables
        if table.table_id == "funding_rank_trx"
    )
    assert len(table.rows) == 4
    assert "Top 4" in table.title
    assert "0.64" not in str(table.rows)


def test_63_fixed_amount_candidates_are_not_formally_ranked():
    display = prepare_report_for_display(_document())
    section = _section(display, "transfer_patterns")
    assert "未保存各值出現次數及占比" in " ".join(section.content_blocks)
    assert "Observation ID" not in str(section.tables)
    assert "主要固定金額" not in str(section.tables)


def test_64_deterministic_toc_is_generated_from_existing_sections_only():
    display = prepare_report_for_display(_document())
    toc = _section(display, "table_of_contents")
    rendered = " ".join(toc.content_blocks)
    assert "AI 專業綜合" not in rendered
    assert "資料完整度與分析母體" in rendered


def test_65_operation_stage_uses_neutral_names_and_missing_data_wording():
    document = _document()
    sections = tuple(
        replace(
            section,
            tables=tuple(
                replace(
                    table,
                    rows=(
                        table.rows[0],
                        (
                            "來源多元化",
                            table.rows[0][1],
                            table.rows[0][2],
                            "575",
                            *table.rows[0][4:],
                        ),
                    ),
                )
                if table.table_id == "operation_stages"
                else table
                for table in section.tables
            ),
        )
        for section in document.sections
    )
    stages = _section(
        prepare_report_for_display(replace(document, sections=sections)),
        "operation_stages",
    )
    rendered = str(stages.tables)
    assert "初始活動期" in rendered
    assert "後續活動期" in rendered
    assert "來源多元化" not in rendered
    assert "無法判定交易頻率變化" in rendered
    assert "無法判定金額變化" in rendered


def test_66_cover_and_body_do_not_expose_report_enums():
    display = prepare_report_for_display(_document())
    rendered = str(
        tuple(
            (section.title, section.content_blocks, section.tables)
            for section in display.sections
        )
    )
    cover = _section(display, "cover")
    assert "報告類型：確定性分析報告" in cover.content_blocks
    assert "鏈別：TRON" in cover.content_blocks
    assert "deterministic" not in rendered
    assert "Candidate" not in rendered
    assert "Confirmed" not in rendered


def test_67_deterministic_conclusion_does_not_claim_ai_content():
    display = prepare_report_for_display(_document())
    conclusion = _section(display, "conclusion")
    assert "AI 候選解釋" not in " ".join(conclusion.content_blocks)
    assert "規則式觀察與候選解釋" in " ".join(conclusion.content_blocks)


def test_68_provider_complete_follow_up_does_not_claim_provider_gap():
    display = prepare_report_for_display(_document())
    follow_up = _section(display, "recommended_follow_up")
    rendered = " ".join(follow_up.content_blocks)
    assert "Provider 缺漏" not in rendered
    assert "未標記地址身分" in rendered


def test_69_key_address_table_precedes_asset_overview_and_analysis():
    display = prepare_report_for_display(_two_asset_document())
    ids = [section.section_id for section in display.sections]
    assert ids.index("key_addresses") < ids.index("asset_flows")
    assert ids.index("key_addresses") < ids.index("asset_analysis_usdt")


def test_70_key_address_table_is_complete_and_trace_prioritized():
    display = prepare_report_for_display(_two_asset_document())
    table = next(
        table
        for table in _section(display, "key_addresses").tables
        if table.table_id == "key_address_summary"
    )
    assert table.columns == (
        "調查角色",
        "完整地址（地址編號）",
        "資產",
        "流入／流出金額",
        "追蹤優先級",
    )
    assert any("調查標的" in row[0] and ADDRESS in row[1] for row in table.rows)
    assert any("USDT" in row[0] for row in table.rows)
    assert all("…" not in row[1] for row in table.rows)


def test_71_report_is_explicitly_first_hop_not_off_ramp_confirmation():
    display = prepare_report_for_display(_two_asset_document())
    rendered = str(
        tuple(
            (section.title, section.content_blocks, section.tables)
            for section in display.sections
        )
    )
    assert display.title == "地址剖繪與第一層資金流分析報告"
    assert "Address Profile and First-Hop Fund Flow Analysis" in rendered
    assert "尚未完成 transaction-level path tracing" in rendered
    assert "不據此確認最終下車點或資金最終受益人" in rendered
    assert "來源與去向並列僅為排名關聯摘要" in rendered
    assert "已確認最終下車點" not in rendered
    assert "已完成多層追蹤" not in rendered


def test_72_principal_value_addresses_precede_operational_addresses():
    table = next(
        table
        for table in _section(
            prepare_report_for_display(_two_asset_document()),
            "key_addresses",
        ).tables
        if table.table_id == "key_address_summary"
    )
    usdt_index = next(
        index for index, row in enumerate(table.rows) if row[2] == "USDT"
    )
    trx_index = next(
        index for index, row in enumerate(table.rows) if row[2] == "TRX"
    )
    assert usdt_index < trx_index
    assert table.rows[trx_index][4] == "營運型"


def test_73_core_address_table_is_bounded_and_not_address_id_sorted():
    table = next(
        table
        for table in _section(
            prepare_report_for_display(_two_asset_document()),
            "key_addresses",
        ).tables
        if table.table_id == "key_address_summary"
    )
    assert 1 <= len(table.rows) <= 15
    priorities = [row[4] for row in table.rows]
    assert priorities.index("中") < priorities.index("營運型")


def test_74_trx_only_report_is_downgraded_and_operational():
    document = _trx_only_document()
    assert document.metadata.scope_assets == ("TRX",)
    assert document.metadata.principal_assets == ("USDT",)
    assert document.metadata.principal_asset_coverage == "missing"
    assert document.metadata.full_address_profile is False
    assert document.metadata.first_hop_fund_flow_complete is False
    assert document.metadata.off_ramp_analysis_available is False

    display = prepare_report_for_display(document)
    assert display.title == "TRX 子資產分析與交易對手概覽"
    assert (
        _section(display, "asset_analysis_trx").title
        == "TRX 營運資產與費用型對手方分析"
    )
    rendered = str(display.sections)
    assert "不包含本案主要價值資產之完整資金流" in rendered
    assert "不等同主要價值資產下車點" in rendered
    assert "已識別下車點" not in rendered


def test_75_core_address_section_is_before_completeness_and_assets():
    display = prepare_report_for_display(_trx_only_document())
    ids = [
        section.section_id
        for section in display.sections
    ]
    assert ids.index("target") < ids.index("key_addresses")
    assert ids.index("key_addresses") < ids.index("completeness")
    assert ids.index("key_addresses") < ids.index("asset_flows")
    assert _section(display, "key_addresses").title == "核心地址對照表"
